import json
import re
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from copy import deepcopy
from openai import OpenAI
import os
import zipfile
import difflib

try:
    import pdfplumber
except ImportError:
    pdfplumber = None


# ================== 文件格式预检 (保持不变) ==================
def validate_file_format(file_path):
    if not os.path.exists(file_path):
        return False, "文件上传失败，未能保存到服务器。"
    filename = os.path.basename(file_path)
    ext = os.path.splitext(filename)[1].lower()

    if ext == '.docx':
        if not zipfile.is_zipfile(file_path):
            return False, f"❌ 文件【{filename}】格式错误！\n它看起来像是旧版 .doc 或已损坏。\n💡 请用 Word 打开并‘另存为’ .docx 格式。"
        try:
            Document(file_path)
        except Exception as e:
            return False, f"❌ 文件【{filename}】内容损坏: {str(e)}"
    elif ext == '.pdf':
        if pdfplumber is None:
            return False, "缺少 pdfplumber 库。"
        try:
            with pdfplumber.open(file_path) as pdf:
                if len(pdf.pages) == 0: return False, "❌ PDF 文件是空的。"
        except Exception as e:
            return False, f"❌ PDF 损坏: {str(e)}"
    return True, "OK"


# ================= 文本读取 (保持不变) =================
def _read_pdf(file_path):
    if pdfplumber is None: return ""
    text_content = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                txt = page.extract_text()
                if txt: text_content.append(f"[PDF_第{i + 1}页] {txt}")
                tables = page.extract_tables()
                for t_idx, table in enumerate(tables):
                    clean_table = []
                    for row in table:
                        clean_row = [str(c).replace('\n', ' ') for c in row if c]
                        if clean_row: clean_table.append(" | ".join(clean_row))
                    if clean_table:
                        text_content.append(f"[PDF_表格_{i + 1}_{t_idx}]\n" + "\n".join(clean_table))
    except Exception as e:
        return f"[PDF读取失败] {str(e)}"
    return "\n".join(text_content)


def read_file_content(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf': return _read_pdf(file_path)
    try:
        doc = Document(file_path)
        text = []
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_txt = " | ".join([c.text.strip() for c in row.cells if c.text.strip()])
                if row_txt: table_data.append(row_txt)
            if table_data:
                text.append(f"【表格区_{i}】\n" + "\n".join(table_data))

        para_data = []
        for p in doc.paragraphs:
            if p.text.strip(): para_data.append(p.text.strip())
        if para_data:
            text.append("【正文区】\n" + "\n".join(para_data))

        return "\n\n".join(text)
    except Exception as e:
        return f"[读取错误] {str(e)}"


def _count_distinct_cells(row):
    elements = {id(cell._element) for cell in row.cells}
    return len(elements)


def _count_filled_cells(row):
    return sum(1 for cell in row.cells if cell.text.strip())


def _find_keyword_location(doc, keyword):
    for table in doc.tables:
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if keyword and keyword in cell.text:
                    return table, r_idx, c_idx
    return None, -1, -1


def _find_keyword_location_fuzzy(doc, keyword, threshold=0.85):
    if not keyword:
        return None, -1, -1
    best = (None, -1, -1, 0.0)
    for table in doc.tables:
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                score = get_fuzzy_score(keyword, cell.text)
                if score > best[3]:
                    best = (table, r_idx, c_idx, score)
    if best[3] >= threshold:
        return best[0], best[1], best[2]
    return None, -1, -1


def _table_distinct_cols(table, max_rows=6):
    distinct_counts = []
    for row in table.rows[:max_rows]:
        distinct_counts.append(_count_distinct_cells(row))
    return max(distinct_counts) if distinct_counts else 0


def _row_header_hit_count(row, headers):
    if not row or not headers:
        return 0
    row_text = "".join(cell.text for cell in row.cells)
    return sum(1 for h in headers if h and h in row_text)


def _cell_vertically_merged(table, row_idx, col_idx):
    if row_idx + 1 >= len(table.rows):
        return False
    return table.rows[row_idx].cells[col_idx]._element is table.rows[row_idx + 1].cells[col_idx]._element


def normalize_plan_with_template(plan, template_path):
    """
    根据模板结构，将误判的 lists 自动降级为 kv，避免破坏格式。
    """
    if not plan or not plan.get("lists"):
        return plan
    if not zipfile.is_zipfile(template_path):
        return plan

    doc = Document(template_path)
    kv = plan.get("kv", [])
    lists = []

    for item in plan.get("lists", []):
        keyword = item.get("keyword", "")
        headers = item.get("headers", [])
        data = item.get("data", [])

        table, r_idx, c_idx = _find_keyword_location(doc, keyword)
        if table is None:
            table, r_idx, c_idx = _find_keyword_location_fuzzy(doc, keyword)
        if table is None:
            lists.append(item)
            continue

        row = table.rows[r_idx]
        distinct_cells = _count_distinct_cells(row)
        filled_cells = _count_filled_cells(row)

        next_row = table.rows[r_idx + 1] if r_idx + 1 < len(table.rows) else None
        next_distinct = _count_distinct_cells(next_row) if next_row else 0
        next_filled = _count_filled_cells(next_row) if next_row else 0

        header_hits = max(
            _row_header_hit_count(row, headers),
            _row_header_hit_count(next_row, headers)
        )

        data_width = 0
        for row_data in data:
            if isinstance(row_data, (list, tuple)):
                data_width = max(data_width, len(row_data))
            else:
                data_width = max(data_width, 1)

        distinct_cols = _table_distinct_cols(table)
        is_vertically_merged = _cell_vertically_merged(table, r_idx, c_idx)

        looks_like_table = (
            distinct_cols >= 3
            or next_distinct >= 3
            or (header_hits >= 2 and distinct_cols >= max(2, len(headers)))
            or (filled_cells >= 2 and next_filled >= 2 and distinct_cols >= 2)
        )

        if data_width <= 1 and distinct_cols <= 2:
            looks_like_table = False
        if is_vertically_merged and distinct_cols <= 2 and header_hits == 0:
            looks_like_table = False

        if looks_like_table:
            lists.append(item)
            continue

        row_lines = []
        for row_data in data:
            if isinstance(row_data, (list, tuple)):
                line = " ".join(str(cell).strip() for cell in row_data if str(cell).strip())
            else:
                line = str(row_data).strip()
            if line:
                row_lines.append(line)
        merged_text = "\n".join(row_lines)

        if merged_text:
            existing = next((entry for entry in kv if entry.get("anchor") == keyword), None)
            if existing:
                if existing.get("val"):
                    existing["val"] = f"{existing['val']}\n{merged_text}"
                else:
                    existing["val"] = merged_text
            else:
                kv.append({"anchor": keyword, "val": merged_text})

    plan["kv"] = kv
    plan["lists"] = lists
    return plan


def extract_docx_preview(file_path, max_paragraphs=20, max_tables=5, max_rows=8):
    if not zipfile.is_zipfile(file_path):
        return {"paragraphs": [], "tables": []}
    doc = Document(file_path)

    paragraphs = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            paragraphs.append(text)
        if len(paragraphs) >= max_paragraphs:
            break

    tables = []
    for table in doc.tables[:max_tables]:
        rows = []
        for row in table.rows[:max_rows]:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(rows)

    return {"paragraphs": paragraphs, "tables": tables}


# ================= V5 核心 Prompt (修复基础信息遗漏) =================
def generate_filling_plan_v2(client, old_data, target_structure, model="deepseek-chat", temperature=0.25,
                             max_tokens=None, return_usage=False):
    prompt = f"""
    你是一个专业的数据迁移专家。

    【源数据】
    {old_data[:12000]} 

    【目标表结构】
    {target_structure[:4000]}

    【必须严格执行的指令】
    1. **全面提取 KV (基础信息 + 软信息)**:
       - **基础信息**: 必须地毯式提取所有短字段！包括“学号”、“性别”、“民族”、“籍贯”、“政治面貌”、“出生年月”等。不要因为它们简单就忽略！
       - **软信息**: 对于“自我鉴定”、“主要事迹”等长文本，如果源数据没有，请**根据简历事实自动撰写**，禁止留空。

    2. **Lists (多行表格)**:
       - 凡是目标表中有明确表头（如：时间|课程|成绩）的，必须提取为 `lists`。
       - **严格对齐**: `headers` 列数必须与 `data` 列数一致。

    3. **Checkbox (勾选框)**:
       - 寻找“□”符号。
       - 输出 keyword (选项文字) 和 status (有/无/是/否)。

    【输出格式 (JSON)】
    {{
        "kv": [
            {{"anchor": "姓名", "val": "张三"}},
            {{"anchor": "学号", "val": "20201101"}},
            {{"anchor": "性别", "val": "男"}},
            {{"anchor": "自我鉴定", "val": "本人在校期间..."}}
        ],
        "checkbox": [
            {{"keyword": "党员", "status": "有"}},
            {{"keyword": "英语六级", "status": "无"}}
        ],
        "lists": [
            {{
                "keyword": "获奖情况", 
                "headers": ["时间", "奖项", "等级"],
                "data": [["2023.09", "一等奖", "校级"]]
            }}
        ]
    }}
    """
    request_kwargs = {
        "model": model,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": temperature
    }
    if max_tokens is not None:
        request_kwargs["max_tokens"] = max_tokens
    response = client.chat.completions.create(**request_kwargs)
    content = response.choices[0].message.content
    content = re.sub(r'```json\s*|\s*```', '', content)

    try:
        plan = json.loads(content)

        # 自动清洗表格列数 (防止报错)
        if "lists" in plan:
            for lst in plan["lists"]:
                headers = lst.get("headers", [])
                data = lst.get("data", [])
                if headers and data:
                    num_cols = len(headers)
                    cleaned_data = []
                    for row in data:
                        if len(row) > num_cols:
                            cleaned_data.append(row[:num_cols])
                        elif len(row) < num_cols:
                            cleaned_data.append(row + [""] * (num_cols - len(row)))
                        else:
                            cleaned_data.append(row)
                    lst["data"] = cleaned_data

        if return_usage:
            return plan, getattr(response, "usage", None)
        return plan
    except:
        if return_usage:
            return {"kv": [], "checkbox": [], "lists": []}, getattr(response, "usage", None)
        return {"kv": [], "checkbox": [], "lists": []}


def refine_text_v2(client, original_text, instruction, model="deepseek-chat", temperature=0.7, max_tokens=None,
                   return_usage=False):
    prompt = f"原文：{original_text}\n指令：{instruction}\n请输出修改后的结果："
    request_kwargs = {
        "model": model,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": temperature
    }
    if max_tokens is not None:
        request_kwargs["max_tokens"] = max_tokens
    response = client.chat.completions.create(**request_kwargs)
    content = response.choices[0].message.content
    if return_usage:
        return content, getattr(response, "usage", None)
    return content


# ================= 写入逻辑 =================

def force_write_cell(cell, text, alignment="auto"):
    """
    格式美化：清除原有格式，自动判断居中或左对齐
    """
    # 保护性检查：如果 text 是 None，转为空字符串
    if text is None: text = ""

    cell._element.clear_content()
    p = cell.add_paragraph()

    text_len = len(str(text))
    if alignment == "auto":
        # 短文本居中，长文本左对齐
        if text_len < 15 and "\n" not in str(text):
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    run = p.add_run(str(text))
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0, 0, 0)


def append_write_cell(cell, text, alignment="left"):
    """
    保留原有内容，追加新内容（避免覆盖标题/标签）
    """
    if text is None:
        text = ""
    if not cell.text.strip():
        force_write_cell(cell, text, alignment=alignment)
        return

    p = cell.add_paragraph()
    if alignment == "left":
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p.add_run(str(text))
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0, 0, 0)


def get_next_distinct_cell(row, current_idx):
    current_cell = row.cells[current_idx]
    for i in range(current_idx + 1, len(row.cells)):
        next_cell = row.cells[i]
        if next_cell._element is not current_cell._element:
            return next_cell
    return None


def find_next_writable_cell(row, current_idx):
    current_cell = row.cells[current_idx]
    for i in range(current_idx + 1, len(row.cells)):
        next_cell = row.cells[i]
        if next_cell._element is current_cell._element:
            continue
        if not next_cell.text.strip():
            return next_cell
    return None


def handle_checkbox(cell, status):
    text = cell.text
    new_text = text

    # ====== 【关键修复】: 使用穷举匹配法，解决“有□ 无□”同格问题 ======

    # 1. 映射表：状态 -> 需要替换的目标字符串
    replace_map = {}

    if status in ["无", "No", "否", "None", "未通过"]:
        replace_map = {
            "无□": "无☑", "无 □": "无 ☑",
            "否□": "否☑", "否 □": "否 ☑",
            "未通过□": "未通过☑"
        }
    elif status in ["有", "Yes", "是", "Have", "通过"]:
        replace_map = {
            "有□": "有☑", "有 □": "有 ☑",
            "是□": "是☑", "是 □": "是 ☑",
            "通过□": "通过☑"
        }

    # 2. 优先尝试包含文字的精确替换 (解决 "有□ 无□" 这种场景)
    replaced_flag = False
    for k, v in replace_map.items():
        if k in text:
            new_text = new_text.replace(k, v)
            replaced_flag = True

    # 3. 如果没找到带文字的框，但确实是勾选状态，且格子里只有一个框，则兜底替换
    if not replaced_flag and "□" in text:
        # 只有当状态明确为肯定，或者明确针对该项时才打钩
        if status in ["有", "Yes", "是", "Have", "True"]:
            new_text = text.replace("□", "☑", 1)  # 只替第一个，防止误伤

    if new_text != text:
        cell.text = new_text
        return True
    return False


def deepcopy_row(table, source_row):
    tbl = table._tbl
    tr = deepcopy(source_row._tr)
    source_row._tr.addprevious(tr)
    return table.rows[source_row._index - 1]


def get_fuzzy_score(anchor, target_text):
    a = anchor.replace(" ", "").replace("\n", "").lower()
    t = target_text.replace(" ", "").replace("\n", "").lower()
    if not a or not t: return 0.0
    if a == t: return 1.0
    if a in t: return 1.0
    return difflib.SequenceMatcher(None, a, t).ratio()


def is_placeholder_text(text):
    clean = text.replace(" ", "")
    return any(k in clean for k in ["此栏", "填写", "说明", "简述", "备注"])


# --- 辅助函数：设置单元格为纵向合并的“继续”状态 ---
def set_cell_merge_continue(cell):
    """
    修改单元格 XML，使其属性变为 vMerge="continue" (即合并单元格的非首行部分)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # 查找现有的 vMerge
    vMerge = tcPr.find(qn('w:vMerge'))
    if vMerge is None:
        vMerge = OxmlElement('w:vMerge')
        tcPr.append(vMerge)
    # vMerge 标签没有 val 属性时，默认为 continue
    if 'w:val' in vMerge.attrib:
        del vMerge.attrib['w:val']

# --- 辅助函数：在指定行之后插入新行 ---
def insert_row_after(table, ref_row):
    """
    在 ref_row 之后插入一行，并复制 ref_row 的样式
    """
    tbl = table._tbl
    new_tr = deepcopy(ref_row._tr)
    ref_row._tr.addnext(new_tr)
    # 找到新插入的行对象
    new_row_idx = ref_row._index + 1
    return table.rows[new_row_idx]

def get_row_merge_range(table, row_idx, col_idx):
    """
    计算纵向合并范围 (start, end)
    """
    start_cell = table.rows[row_idx].cells[col_idx]
    start_element = start_cell._element
    end_row = row_idx
    for r in range(row_idx + 1, len(table.rows)):
        current_cell = table.rows[r].cells[col_idx]
        if current_cell._element is start_element:
            end_row = r
        else:
            break
    return row_idx, end_row

def find_column_index_by_header(row, header_texts):
    mapping = {}
    for idx, cell in enumerate(row.cells):
        txt = cell.text.strip().replace(" ", "")
        for h in header_texts:
            if h in txt:
                mapping[h] = idx
    return mapping

def execute_word_writing_v2(plan, template_path, output_path, progress_callback=None):
    if not zipfile.is_zipfile(template_path):
        raise ValueError("目标文件格式错误")
    doc = Document(template_path)

    # ---------------- 1. KV 写入 ----------------
    total_kv = len(plan.get("kv", []))
    anchor_usage = {}
    for i, item in enumerate(plan.get("kv", [])):
        anchor, val = item["anchor"], item["val"]
        if not val: continue

        if progress_callback: progress_callback(int(10 + (i / total_kv) * 30), f"正在写入: {anchor}...")

        found = False
        for table in doc.tables:
            for row in table.rows:
                for c_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip().replace(" ", "")
                    clean_anchor = anchor.strip().replace(" ", "")
                    match_score = get_fuzzy_score(clean_anchor, cell_text)
                    strict_match = clean_anchor in cell_text or cell_text in clean_anchor
                    short_anchor = len(clean_anchor) <= 8
                    if (short_anchor and not strict_match):
                        continue

                    if strict_match or match_score > 0.9:
                        used_cells = anchor_usage.get(anchor, set())
                        if cell._element in used_cells:
                            continue
                        target_cell = None
                        candidate = find_next_writable_cell(row, c_idx)

                        # 优先写入相邻可写单元格
                        if candidate:
                            target_cell = candidate
                        else:
                            # 无相邻单元格时，才考虑写入当前单元格
                            target_cell = cell

                        if target_cell:
                            # 保护机制：防止覆盖表头
                            # 如果目标格子很短，且包含冒号或看起来像另一个表头，跳过
                            if len(target_cell.text) < 10 and ("：" in target_cell.text or ":" in target_cell.text):
                                pass
                            else:
                                if target_cell is cell:
                                    append_write_cell(target_cell, val, alignment="left")
                                else:
                                    force_write_cell(target_cell, val, alignment="auto")
                                used_cells.add(cell._element)
                                anchor_usage[anchor] = used_cells
                                found = True
                                break
                if found: break
            if found: break

    # ---------------- 2. Checkbox 写入 (新版匹配逻辑) ----------------
    if progress_callback: progress_callback(60, "处理勾选框...")
    for item in plan.get("checkbox", []):
        keyword, status = item["keyword"], item["status"]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # 只有当关键字匹配时才尝试打钩
                    if keyword in cell.text:
                        handle_checkbox(cell, status)

    # 3. Lists 写入 (✨ 修复表头被顶飞的问题 ✨)
    if progress_callback: progress_callback(80, "处理表格列表...")

    for item in plan.get("lists", []):
        keyword = item["keyword"]
        headers = item.get("headers", [])
        data = item.get("data", [])

        if not data: continue

        # A. 定位锚点
        target_table = None
        anchor_row_idx = -1
        anchor_col_idx = -1

        found = False
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    if keyword in cell.text:
                        target_table = table
                        anchor_row_idx = r_idx
                        anchor_col_idx = c_idx
                        found = True
                        break
                if found: break
            if found: break

        if not found:
            continue

        # B. 判读当前版块类型
        start_r, end_r = get_row_merge_range(target_table, anchor_row_idx, anchor_col_idx)
        is_side_block = (end_r > start_r)  # 是否为侧边栏合并类型

        # C. 智能确定数据起始行 (Fix: 防止写在表头上面)
        header_map = find_column_index_by_header(target_table.rows[anchor_row_idx], headers)
        data_start_row = anchor_row_idx  # 默认从锚点行开始算

        # 策略：向下一行探测
        if anchor_row_idx + 1 < len(target_table.rows):
            next_row = target_table.rows[anchor_row_idx + 1]
            next_row_text = "".join([c.text for c in next_row.cells]).strip()

            # 1. 尝试在下一行精准匹配表头
            candidate_map = find_column_index_by_header(next_row, headers)

            if candidate_map:
                # 命中表头！
                header_map = candidate_map
                data_start_row = anchor_row_idx + 1

            elif not is_side_block:
                # 2. 【关键修复】如果是普通模式，且没匹配到表头，但下一行明显有文字
                # 我们假设下一行就是表头（比如 Word 里是“课程名称”，AI 识别成“课程名”，导致没匹配上）
                # 这种情况下，我们强制跳过下一行，从下下行开始写
                if len(next_row_text) > 2 and (
                        "课程" in next_row_text or "名称" in next_row_text or "成绩" in next_row_text or "Date" in next_row_text):
                    data_start_row = anchor_row_idx + 1  # 把下一行视为表头

        # 确定游标初始位置：总是从 data_start_row 的下一行开始写
        cursor_row_idx = data_start_row + 1

        # 【关键修复】同步边界 end_r
        # 如果是普通模式，end_r 初始值可能就是锚点行(0)。
        # 但如果我们要从 Row 2 开始写，必须把 end_r 至少推到 Row 2 的前一行，防止 cursor > end_r 导致在 Row 0 后面插入
        # 简单来说：在普通模式下，只要表格里还有空行，就不要急着插入。
        if not is_side_block:
            # 只要 cursor 指向的行存在，我们就认为它在边界内
            if cursor_row_idx < len(target_table.rows):
                end_r = max(end_r, cursor_row_idx)

        # D. 循环填入数据
        for data_idx, data_row in enumerate(data):

            # 检查是否越界/需要扩容
            if cursor_row_idx > end_r:

                # === 扩容逻辑 ===
                # 复制模板：普通模式下，尽量复制上一行（即数据行样式）；侧边栏模式复制最后一行
                template_row_idx = end_r
                if not is_side_block and cursor_row_idx > 0:
                    template_row_idx = cursor_row_idx - 1

                # 安全检查
                if template_row_idx >= len(target_table.rows): template_row_idx = len(target_table.rows) - 1

                last_row = target_table.rows[template_row_idx]
                new_row = insert_row_after(target_table, last_row)

                # === 样式处理 ===
                if is_side_block:
                    # 侧边栏模式：保留锚点列，清空其他
                    for idx, cell in enumerate(new_row.cells):
                        if idx == anchor_col_idx:
                            pass
                        else:
                            cell._element.clear_content()
                    # 修复左侧合并
                    merge_cell = new_row.cells[anchor_col_idx]
                    set_cell_merge_continue(merge_cell)
                else:
                    # 普通模式：清空所有列，不合并
                    for cell in new_row.cells:
                        cell._element.clear_content()

                end_r += 1
                # ===================

            # E. 执行写入
            if cursor_row_idx >= len(target_table.rows): break
            current_row = target_table.rows[cursor_row_idx]

            if header_map:
                # 有表头映射
                for h_text, col_idx in header_map.items():
                    try:
                        val_idx = headers.index(h_text)
                        if val_idx < len(data_row):
                            force_write_cell(current_row.cells[col_idx], data_row[val_idx])
                    except:
                        pass
            else:
                # 无表头映射（盲填）
                start_col = anchor_col_idx + 1 if is_side_block else 0
                write_col = start_col
                data_ptr = 0
                while write_col < len(current_row.cells) and data_ptr < len(data_row):
                    cell = current_row.cells[write_col]
                    # 跳过合并列(水平)
                    if write_col > 0 and cell._element is current_row.cells[write_col - 1]._element:
                        write_col += 1
                        continue
                    force_write_cell(cell, data_row[data_ptr])
                    data_ptr += 1
                    write_col += 1

            cursor_row_idx += 1

    doc.save(output_path)
    if progress_callback: progress_callback(100, "完成")
